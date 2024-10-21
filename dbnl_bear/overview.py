import docx
from docx.shared import RGBColor
from fuzzywuzzy import fuzz
import matplotlib.pyplot as plt
import tqdm 

""""
The highlight relevant passages function still has flaws. 
I have not been able to come up with a system that has 100% recall and 100% precision,
but theoretically that should be possible.

The barcode plot has the same issue (with a different solution though) but here the problem
is less urgent because the barcode plot is just used to give an indication so there's more
room for error.
"""
import matplotlib.pyplot as plt
import docx
from difflib import SequenceMatcher
import re
import tqdm
from fuzzywuzzy import fuzz
import numpy as np

def highlight_relevant_passages(input_file, output_file, relevant_passages):
    # Read the original text
    with open(input_file, 'r', encoding='utf-8') as file:
        original_text = file.read()

    # Create a new Word document
    doc = docx.Document()

    # Preprocess the text to create an index
    text_index = {}
    for i in range(len(original_text) - 5):
        five_gram = original_text[i:i+5]
        if five_gram not in text_index:
            text_index[five_gram] = []
        text_index[five_gram].append(i)

    # Process the text and add it to the document with highlighting
    current_pos = 0
    for passage in tqdm.tqdm(relevant_passages, desc="Processing passages"):
        # Find the best match for the passage using the index
        best_match = None
        best_ratio = 0
        potential_starts = text_index.get(passage[:5], [])
        
        for start in potential_starts:
            end = start + len(passage)
            substring = original_text[start:end]
            ratio = fuzz.ratio(passage, substring)
            if ratio > best_ratio:
                best_ratio = ratio
                best_match = (start, end)

        if best_match:
            doc.add_paragraph(original_text[current_pos:best_match[0]])

            # Add the matched text with yellow highlighting
            p = doc.add_paragraph()
            run = p.add_run(original_text[best_match[0]:best_match[1]])
            run.font.highlight_color = docx.enum.text.WD_COLOR_INDEX.YELLOW

            current_pos = best_match[1]

    # Add any remaining text
    if current_pos < len(original_text):
        doc.add_paragraph(original_text[current_pos:])

    # Save the document
    doc.save(output_file)

def map_to_original_index(original_text, alphanumeric_index):
    alphanumeric_count = 0
    for i, char in enumerate(original_text):
        if char.isalnum():
            if alphanumeric_count == alphanumeric_index:
                return i
            alphanumeric_count += 1
    return len(original_text)

def create_plain_text_output(output_file, document_name, relevant_passages):
    with open(output_file, 'w', encoding='utf-8') as file:
        file.write(f"In this text: {document_name} the following has been found (numbered):\n")
        for n, passage in enumerate(relevant_passages):
            file.write(f"No. {n}:    {passage}\n")

def create_barcode_plot(relevant_passages, all_passages):
    barcode = np.array([1 if p in relevant_passages else 0 for p in all_passages])
    pixel_per_bar = 4
    dpi = 100
    fig = plt.figure(figsize=(len(barcode) * pixel_per_bar / dpi, 2), dpi=dpi)
    ax = fig.add_axes([0, 0, 1, 1])  # span the whole figure
    ax.set_axis_off()
    ax.imshow(barcode.reshape(1, -1), cmap='binary', aspect='auto',
          interpolation='nearest')
    plt.show()
